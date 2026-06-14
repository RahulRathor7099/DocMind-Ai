import os
import shutil
from PIL import Image, ImageDraw

def generate_txt_file(path, content):
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Generated text file: {path}")

def generate_docx_file(path, title, paragraphs, table_data=None):
    import docx
    doc = docx.Document()
    doc.add_heading(title, level=0)
    
    for p in paragraphs:
        doc.add_paragraph(p)
        
    if table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        table = doc.add_table(rows=rows, cols=cols)
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = str(val)
                
    doc.save(path)
    print(f"Generated docx file: {path}")

def generate_image_file(path, text_lines, is_pdf=False):
    # Create white canvas
    width, height = 800, 1100
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    # Draw simple lines of text using default font
    y_offset = 50
    for line in text_lines:
        d.text((50, y_offset), line, fill=(0, 0, 0))
        y_offset += 30
        
    if is_pdf:
        img.save(path, "PDF")
        print(f"Generated PDF from image canvas: {path}")
    else:
        img.save(path, "PNG")
        print(f"Generated PNG image: {path}")

def main():
    # Setup directory paths relative to this script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    target_dir = os.path.join(project_root, "sample_documents")
    os.makedirs(target_dir, exist_ok=True)
    
    # 1. Sample 1: Financial Report (PDF with Embedded Tables)
    # Attempt to copy the existing test PDF if it exists, otherwise generate a new one
    pdf_target = os.path.join(target_dir, "sample1_financial_report.pdf")
    pdf_source = os.path.join(script_dir, "uploads", "documents", "10", "688704bcc9e0460e8ec10fbac6f8edd6.pdf")
    
    if os.path.exists(pdf_source):
        shutil.copy(pdf_source, pdf_target)
        print(f"Copied existing test PDF to: {pdf_target}")
    else:
        pdf_text = [
            "DocMind AI - Quarterly Financial Performance Report",
            "Report Type: Financial Statement",
            "Topic: Q3 Financial Summary",
            "Sensitivity: Internal",
            "",
            "Section 1: Executive Summary",
            "OpenVision Technologies Pvt Ltd is showing exceptional growth in AI services.",
            "Revenue for Q3 has increased by 30% compared to Q2, driven by new enterprise contract acquisitions.",
            "Operational costs remain within budget limits, leading to an improved profit margin.",
            "",
            "Section 2: Key Financial Metrics",
            "Below is the summary of quarterly performance for the fiscal year 2026.",
            "Please refer to the table for detailed revenue breakdown.",
            "",
            "Quarter  | Revenue  | Expenses | Net Profit",
            "Q1       | $120,000 | $80,000  | $40,000",
            "Q2       | $180,000 | $110,000 | $70,000",
            "Q3       | $240,000 | $130,000 | $110,000",
            "",
            "Approved by: Chief Financial Officer"
        ]
        generate_image_file(pdf_target, pdf_text, is_pdf=True)

    # 2. Sample 2: Employee Handbook (TXT file)
    txt_path = os.path.join(target_dir, "sample2_employee_handbook.txt")
    txt_content = """DocMind AI - Employee Conduct Policy Handbook
Sensitivity Level: Internal
Topic: Human Resources & Workplace Standards
Domain: Business

Welcome to DocMind AI. This document outlines the general code of conduct and compliance standards for all team members.

1. Work Hours & Remote Work
Standard work hours are flexible. Core collaboration hours are 10:00 AM to 4:00 PM. Remote work is supported; ensure your remote setup complies with our IT security protocols.

2. Intellectual Property
All code, models, and documentation created during your employment belong exclusively to DocMind AI.

3. Professional Conduct
We expect all employees to maintain a respectful, inclusive, and professional environment. Any form of harassment or discrimination is strictly prohibited.
"""
    generate_txt_file(txt_path, txt_content)

    # 3. Sample 3: Office Lease Agreement (DOCX file)
    docx_path = os.path.join(target_dir, "sample3_lease_agreement.docx")
    docx_paragraphs = [
        "This Commercial Lease Agreement (the 'Lease') is entered into as of January 1, 2026.",
        "LANDLORD: Realty Corp Properties Ltd.",
        "TENANT: DocMind AI Technologies.",
        "PREMISES: Suite 404, Tech Hub Plaza, Sector 62.",
        "1. TERM: The term of this Lease shall be for 24 months, starting January 1, 2026.",
        "2. RENT: Tenant agrees to pay Landlord monthly rent as detailed in the payment schedule below.",
        "3. SECURITY DEPOSIT: Tenant shall deposit the sum of $10,000 as security for performance under this lease."
    ]
    docx_table = [
        ["Month Range", "Monthly Rent", "Due Date"],
        ["Months 1-12", "$3,500", "1st of every month"],
        ["Months 13-24", "$3,800", "1st of every month"]
    ]
    generate_docx_file(
        docx_path, 
        "Commercial Lease Agreement", 
        docx_paragraphs, 
        table_data=docx_table
    )

    # 4. Sample 4: Scanned Receipt (PNG file for OCR testing)
    receipt_path = os.path.join(target_dir, "sample4_scanned_receipt.png")
    receipt_text = [
        "SUPERMARKET OUTLET #120",
        "Date: 14-06-2026",
        "Time: 14:32:00",
        "",
        "ITEMS PURCHASED:",
        "------------------------------------",
        "1. Laser Printer       $150.00",
        "2. Printer Paper A4     $20.00",
        "3. USB-C Cable (2m)     $15.00",
        "4. Wireless Keyboard    $45.00",
        "------------------------------------",
        "SUBTOTAL:              $230.00",
        "TAX (8.5%):             $19.55",
        "TOTAL AMOUNT PAID:     $249.55",
        "------------------------------------",
        "Payment Method: VISA CREDIT CARD",
        "Card Ending in: 4829",
        "Status: Transaction Approved",
        "",
        "Thank you for shopping with us!"
    ]
    generate_image_file(receipt_path, receipt_text, is_pdf=False)

    # 5. Sample 5: Handwritten Notes (PNG file for OCR testing)
    notes_path = os.path.join(target_dir, "sample5_handwritten_notes.png")
    notes_text = [
        "DocMind AI - Development Brainstorming Session",
        "Topic: Agentic RAG and Table Parsing Ideas",
        "Date: June 12, 2026",
        "",
        "Notes from discussions:",
        "1. For RAG: We should use sentence-transformers for local vector search.",
        "   FAISS is super fast and works offline. No API call costs!",
        "2. Page-level citations are key. User needs to see page thumbnail previews.",
        "   Let's render each page to a PNG using pdf2image during parsing.",
        "3. Classification must run after parsing. Use LLM to extract JSON metadata:",
        "   - document_type (invoice, contract, report, etc.)",
        "   - topic, sensitivity_level, domain, summary",
        "4. Security: sanitize all incoming files, restrict file sizes to 50MB,",
        "   and validate JWT token on every API call.",
        "",
        "Todo list:",
        "- Setup slowapi for rate limiting.",
        "- Build Voice input using browser speech-to-text."
    ]
    generate_image_file(notes_path, notes_text, is_pdf=False)

    # 6. Sample 6: Confidential Policy (TXT file for sensitivity testing)
    conf_path = os.path.join(target_dir, "sample6_confidential_policy.txt")
    conf_content = """PROJECT AURORA - PRODUCT SPECIFICATION DOCUMENT
CLASSIFICATION: CONFIDENTIAL - STRICTLY RESTRICTED TO PROJECT AURORA TEAM
Topic: Next-Gen Agentic RAG Platform Architecture

This document contains trade secrets and confidential product specifications for Project Aurora.

1. Product Vision
Project Aurora is the codename for DocMind AI's next-generation Multi-Agent Collaborative Knowledge Engine. It allows multiple AI subagents to collaborate on resolving user queries with zero latency.

2. Core Architecture
- Front-End Gateway: Real-time WebSockets communication with active status push.
- Embedding Layer: Multi-modal embeddings combining text, layout, and visual features.
- Vector database: Distributed Qdrant cluster with real-time replication.

3. Timeline
- Alpha Release: October 2026
- Beta Testing: December 2026
- Public Launch: February 2027
"""
    generate_txt_file(conf_path, conf_content)

    print("\nSuccessfully generated all sample documents in: " + target_dir)

if __name__ == "__main__":
    main()
