"""
DocMind AI - DOCX Parser
Extracts text from Microsoft Word documents using python-docx
"""

import os
from docx import Document as DocxDocument
from utils.logger import get_logger

logger = get_logger(__name__)


def parse_docx(file_path: str, filename: str, document_id: int) -> dict:
    """
    Parse a .docx file extracting text, tables, and headings.
    
    Args:
        file_path: Path to the .docx file
        filename: Original filename
        document_id: Database document ID
        
    Returns:
        Unified document dict with pages
    """
    try:
        doc = DocxDocument(file_path)
        full_text_parts = []
        tables_data = []

        # Extract paragraphs and headings
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style.name.startswith("Heading"):
                    full_text_parts.append(f"\n## {text}\n")
                else:
                    full_text_parts.append(text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text.strip() for cell in row.cells])

            if table_rows:
                tables_data.append({
                    "table_index": table_idx,
                    "rows": table_rows,
                    "row_count": len(table_rows),
                    "col_count": len(table_rows[0]) if table_rows else 0,
                })

                # Add table text to content
                for row in table_rows:
                    full_text_parts.append(" | ".join(row))

        full_text = "\n".join(full_text_parts)

        # Split into virtual pages (~2000 chars each)
        page_size = 2000
        pages = []
        text_chunks = [full_text[i:i + page_size] for i in range(0, max(1, len(full_text)), page_size)]

        for i, chunk in enumerate(text_chunks):
            if chunk.strip():
                pages.append({
                    "page_number": i + 1,
                    "extracted_text": chunk.strip(),
                    "page_image_path": None,
                    "tables": tables_data if i == 0 else [],
                    "ocr_applied": False,
                })

        logger.info(f"Parsed DOCX '{filename}': {len(pages)} virtual pages, {len(tables_data)} tables")

        return {
            "document_name": filename,
            "total_pages": len(pages),
            "pages": pages,
        }

    except Exception as e:
        logger.error(f"DOCX parsing failed for '{filename}': {e}")
        return {"document_name": filename, "total_pages": 0, "pages": []}
